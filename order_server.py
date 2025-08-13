import os
import platform
from typing import List
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Mm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    import win32api
    import win32print
except ImportError:
    win32api = None
    win32print = None

app = FastAPI(title="Wempy Order & Print Server")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

BASE_DIR = os.path.dirname(__file__)
ORDERS_DIR = os.path.join(BASE_DIR, "orders")
ORDER_ID_FILE = os.path.join(ORDERS_DIR, "last_id.txt")
os.makedirs(ORDERS_DIR, exist_ok=True)

# --- Pydantic Models ---
class OrderItem(BaseModel):
    id: str
    name: str
    qty: int
    price: float # Changed from unitPrice

class Customer(BaseModel):
    name: str
    phone: str
    address: str
    notes: str

class Totals(BaseModel):
    subtotal: float
    delivery: float
    total: float # Changed from grandTotal

class OrderRequest(BaseModel):
    items: List[OrderItem]
    customer: Customer
    totals: Totals

# --- DOCX and Printing Logic ---
def set_cell_text(cell, text, bold=False, align=None):
    """
    Helper function to set text in a table cell, with optional bolding and alignment.
    """
    # Ensure the cell has at least one paragraph
    if not cell.paragraphs:
        cell.add_paragraph()
    
    p = cell.paragraphs[0]
    # Clear existing content by removing all runs
    for run in p.runs:
        r = run._r
        p._p.remove(r)

    run = p.add_run(str(text))
    run.bold = bold
    
    # Set alignment if specified
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # If align is None, it will use the paragraph's default alignment

def format_order_docx(order: OrderRequest, order_id: int, filepath: str):
    if Document is None: raise ImportError("python-docx is not installed")
    doc = Document()

    section = doc.sections[0]
    section.page_width = Mm(72.1)
    section.page_height = Mm(297) # A4 height, can be long
    section.left_margin = Mm(4)
    section.right_margin = Mm(4)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)

    # Set default font and right-to-left alignment for the whole document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- Header ---
    header = doc.add_heading(f"Order #{order_id}", level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p = doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("-" * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Customer Details ---
    doc.add_paragraph().add_run("بيانات العميل").bold = True
    customer_table = doc.add_table(rows=3, cols=2)
    customer_table.style = 'Table Grid'
    # Data on the left, labels on the right
    set_cell_text(customer_table.cell(0, 0), order.customer.name)
    set_cell_text(customer_table.cell(0, 1), "الاسم:", bold=True)
    set_cell_text(customer_table.cell(1, 0), order.customer.phone)
    set_cell_text(customer_table.cell(1, 1), "الهاتف:", bold=True)
    set_cell_text(customer_table.cell(2, 0), order.customer.address)
    set_cell_text(customer_table.cell(2, 1), "العنوان:", bold=True)

    # --- Order Items ---
    doc.add_paragraph().add_run("تفاصيل الطلب").bold = True
    items_table = doc.add_table(rows=1, cols=4)
    items_table.style = 'Table Grid'
    items_table.autofit = False
    items_table.allow_autofit = False

    # NO table_direction = 'rtl'. We simulate it with column order and text alignment.
    # Columns LTR: Total, Price, Qty, Item. Total width ~2.4 inches
    items_table.columns[0].width = Inches(0.6)
    items_table.columns[1].width = Inches(0.5)
    items_table.columns[2].width = Inches(0.2)
    items_table.columns[3].width = Inches(0.8) 

    # Headers are set from left to right to match the visual order
    hdr_cells = items_table.rows[0].cells
    set_cell_text(hdr_cells[0], "الإجمالي", bold=True, align='center')
    set_cell_text(hdr_cells[1], "السعر", bold=True, align='center')
    set_cell_text(hdr_cells[2], "الكمية", bold=True, align='center')
    set_cell_text(hdr_cells[3], "الصنف", bold=True, align='center')

    for item in order.items:
        row_cells = items_table.add_row().cells
        # Data is set from left to right to match the visual order
        set_cell_text(row_cells[0], f"{item.price * item.qty:.2f}")
        set_cell_text(row_cells[1], f"{item.price:.2f}")
        set_cell_text(row_cells[2], str(item.qty))
        set_cell_text(row_cells[3], item.name)

    # --- Totals ---
    doc.add_paragraph().add_run("الحساب").bold = True
    totals_table = doc.add_table(rows=3, cols=2)
    totals_table.style = 'Table Grid'
    set_cell_text(totals_table.cell(0, 0), f"{order.totals.subtotal:.2f} ج.م")
    set_cell_text(totals_table.cell(0, 1), "المجموع الفرعي:", bold=True)
    set_cell_text(totals_table.cell(1, 0), f"{order.totals.delivery:.2f} ج.م")
    set_cell_text(totals_table.cell(1, 1), "رسوم التوصيل:", bold=True)
    set_cell_text(totals_table.cell(2, 0), f"{order.totals.total:.2f} ج.م", bold=True)
    set_cell_text(totals_table.cell(2, 1), "الإجمالي النهائي:", bold=True)

    # --- Notes ---
    if order.customer.notes:
        doc.add_paragraph().add_run("الملاحظات").bold = True
        notes_table = doc.add_table(rows=1, cols=1)
        notes_table.style = 'Table Grid'
        set_cell_text(notes_table.cell(0, 0), order.customer.notes, align='left')

    doc.save(filepath)

def print_file(filepath: str):
    if platform.system() != "Windows":
        print("Printing is only supported on Windows.")
        return
    if not os.path.exists(filepath): raise FileNotFoundError(f"File not found for printing: {filepath}")
    try:
        os.startfile(filepath, "print")
    except Exception as e:
        print(f"os.startfile failed: {e}. Trying win32api.")
        if win32api and win32print:
            try:
                win32api.ShellExecute(0, "print", filepath, f'/d:"{win32print.GetDefaultPrinter()}"', ".", 0)
            except Exception as e2:
                print(f"win32api printing failed: {e2}")

def get_next_order_id() -> int:
    last_id = 0
    if os.path.exists(ORDER_ID_FILE):
        with open(ORDER_ID_FILE, "r") as f:
            content = f.read().strip()
            if content.isdigit(): last_id = int(content)
    new_id = last_id + 1
    with open(ORDER_ID_FILE, "w") as f:
        f.write(str(new_id))
    return new_id

# --- API Endpoints ---
@app.post("/submit_order")
def submit_order(order: OrderRequest):
    try:
        order_id = get_next_order_id()
        docx_path = os.path.join(ORDERS_DIR, f"wempy_order_{order_id}.docx")
        
        format_order_docx(order, order_id, docx_path)
        print_file(docx_path)

        return {"success": True, "order_id": order_id, "file_path": docx_path}
    except Exception as e:
        print(f"Error processing order: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing order: {e}")

@app.get("/print_order/{order_id}")
def reprint_order(order_id: int):
    docx_path = os.path.join(ORDERS_DIR, f"wempy_order_{order_id}.docx")
    if not os.path.exists(docx_path):
        raise HTTPException(status_code=404, detail=f"Order {order_id} not found.")
    try:
        print_file(docx_path)
        return {"success": True, "message": f"Order {order_id} sent to printer."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not print order {order_id}: {e}")

@app.get("/")
def root():
    return {"message": "Wempy Order Service Ready"}

if __name__ == "__main__":
    import uvicorn
    print("Starting Wempy Order & Print Service...")
    uvicorn.run("order_server:app", host="127.0.0.1", port=5000, reload=True)