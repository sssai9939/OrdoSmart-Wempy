import os
import platform
from datetime import datetime
from pathlib import Path
from enums.order_messages import ResponseMessages

try:
    from docx import Document
    from docx.shared import Mm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None
    WD_ALIGN_PARAGRAPH = None  # type: ignore

try:
    import win32api
    import win32print
except ImportError:
    win32api = None  # type: ignore
    win32print = None  # type: ignore

# Project root is the parent of this "app/" directory
ROOT_DIR = Path(__file__).resolve().parent.parent
ORDERS_DIR = ROOT_DIR / "orders"
ORDER_ID_FILE = ORDERS_DIR / "last_id.txt"
ORDERS_DIR.mkdir(exist_ok=True)

def set_cell_text(cell, text, bold=False, align=None):
    """Helper to set text with optional bold and alignment in a python-docx cell."""
    if not getattr(cell, "paragraphs", None):
        cell.add_paragraph()

    p = cell.paragraphs[0]
    for run in list(p.runs):
        r = run._r  # noqa: SLF001
        p._p.remove(r)  # noqa: SLF001

    run = p.add_run(str(text))
    run.bold = bold

    if align and WD_ALIGN_PARAGRAPH:
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'left':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def format_order_docx(order, order_id: int, filepath: Path) -> None:
    if Document is None:
        raise ImportError("python-docx is not installed")

    from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(72.1)
    section.page_height = Mm(297)
    section.left_margin = Mm(4)
    section.right_margin = Mm(4)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    style.paragraph_format.alignment = ALIGN.RIGHT

    header = doc.add_heading(f"Order #{order_id}", level=1)
    header.alignment = ALIGN.CENTER
    date_p = doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_p.alignment = ALIGN.CENTER
    doc.add_paragraph("-" * 30).alignment = ALIGN.CENTER

    doc.add_paragraph().add_run("بيانات العميل").bold = True
    customer_table = doc.add_table(rows=3, cols=2)
    customer_table.style = 'Table Grid'
    set_cell_text(customer_table.cell(0, 0), order.customer.name)
    set_cell_text(customer_table.cell(0, 1), "الاسم", bold=True)
    set_cell_text(customer_table.cell(1, 0), order.customer.phone)
    set_cell_text(customer_table.cell(1, 1), "الهاتف", bold=True)
    set_cell_text(customer_table.cell(2, 0), order.customer.address)
    set_cell_text(customer_table.cell(2, 1), "العنوان", bold=True)
    doc.add_paragraph()

    doc.add_paragraph().add_run("تفاصيل الطلب").bold = True
    items_table = doc.add_table(rows=1, cols=4)
    items_table.style = 'Table Grid'
    items_table.autofit = False
    items_table.allow_autofit = False

    hdr_cells = items_table.rows[0].cells
    set_cell_text(hdr_cells[0], "الإجمالي", bold=True, align='center')
    set_cell_text(hdr_cells[1], "السعر", bold=True, align='center')
    set_cell_text(hdr_cells[2], "الكمية", bold=True, align='center')
    set_cell_text(hdr_cells[3], "الصنف", bold=True, align='center')

    for item in order.items:
        row_cells = items_table.add_row().cells
        set_cell_text(row_cells[0], f"{item.price * item.qty:.2f}")
        set_cell_text(row_cells[1], f"{item.price:.2f}")
        set_cell_text(row_cells[2], str(item.qty))
        set_cell_text(row_cells[3], item.name)

    doc.add_paragraph()

    doc.add_paragraph().add_run("الحساب").bold = True
    totals_table = doc.add_table(rows=3, cols=2)
    totals_table.style = 'Table Grid'
    set_cell_text(totals_table.cell(0, 0), f"{order.totals.subtotal:.2f} ج.م")
    set_cell_text(totals_table.cell(0, 1), "المجموع الفرعي", bold=True)
    set_cell_text(totals_table.cell(1, 0), f"{order.totals.delivery:.2f} ج.م")
    set_cell_text(totals_table.cell(1, 1), "رسوم التوصيل", bold=True)
    set_cell_text(totals_table.cell(2, 0), f"{order.totals.total:.2f} ج.م", bold=True)
    set_cell_text(totals_table.cell(2, 1), "الإجمالي النهائي", bold=True)
    doc.add_paragraph()

    if getattr(order.customer, 'notes', None):
        doc.add_paragraph().add_run("الملاحظات").bold = True
        notes_table = doc.add_table(rows=1, cols=1)
        notes_table.style = 'Table Grid'
        set_cell_text(notes_table.cell(0, 0), order.customer.notes, align='left')
        doc.add_paragraph()

    filepath.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(filepath))

def get_next_order_id() -> int:
    last_id = 0
    if ORDER_ID_FILE.exists():
        content = ORDER_ID_FILE.read_text(encoding="utf-8").strip()
        if content.isdigit():
            last_id = int(content)
    new_id = last_id + 1
    ORDER_ID_FILE.write_text(str(new_id), encoding="utf-8")
    return new_id

def build_order_docx_path(order_id: int) -> Path:
    return ORDERS_DIR / f"wempy_order_{order_id}.docx"

def print_file(filepath: Path) -> None:
    if platform.system() != "Windows":
        print(ResponseMessages.PRINTER_SUPPORTED_WINDOWS.value)
        return

    if not filepath.exists():
        raise FileNotFoundError(f"{ResponseMessages.PRINTER_FILE_NOT_FOUND.value} : {filepath}")

    try:
        os.startfile(str(filepath), "print")  # type: ignore[attr-defined]
    except Exception as e:
        print(f"os.startfile failed: {e}. Trying win32api.")
        if win32api and win32print:
            try:
                win32api.ShellExecute(0, "print", str(filepath), f'/d:"{win32print.GetDefaultPrinter()}"', ".", 0)
            except Exception as e2:
                print(f"win32api printing failed: {e2}")
                